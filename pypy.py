import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from pypdf import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# --- API Key Validation (Replace with your actual logic) ---
def is_valid_api_key(api_key):
    """
    This function should implement your API key validation logic.
    For example, you can check against a list of valid keys,
    perform a database lookup, or make an API call to an authentication service.

    Replace this with your actual validation method.
    """
    # Example: Check against a hardcoded list of valid keys (for testing only)
    valid_keys = ["your_actual_api_key_1", "your_actual_api_key_2"]
    if api_key in valid_keys:
        return True
    else:
        return False
# --- End of API Key Validation ---

# Function to initialize the Gemini model
def initialize_model(api_key):
    genai.configure(api_key=api_key)
    generation_config = {
        "temperature": 0.7,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 8192,
    }
    return genai.GenerativeModel(model_name="gemini-pro", generation_config=generation_config)

# Function to generate questions
def generate_questions(model, topic, num_questions, question_type, show_answers):
    prompt = generate_prompt(topic, num_questions, question_type, show_answers)
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return ""

# Function to generate dynamic prompts with conditional answer generation
def generate_prompt(topic, num_questions, question_type, show_answers):
    if question_type == "Multiple Choice":
        prompt = f"Generate {num_questions} multiple-choice questions on the topic: '{topic}'. " \
                 f"Each question should have 4 options."
        if show_answers:
            prompt += " Indicate the correct answer for each question."
    elif question_type == "Fill in the Blank":
        prompt = f"Generate {num_questions} fill-in-the-blank questions on the topic: '{topic}'."
        if show_answers:
            prompt += " Provide the answer key separately."
    elif question_type == "True or False":
        prompt = f"Generate {num_questions} true or false questions on the topic: '{topic}'."
        if show_answers:
            prompt += " Provide the answer key separately."
    else:  # Default to open-ended questions
        prompt = f"Generate {num_questions} diverse questions on the topic: '{topic}'."
        if show_answers:
            prompt += " Provide detailed answers for each question."
    return prompt

# Function to display questions with answer show/hide option
def display_questions(questions, question_type, show_answers):
    if questions:
        st.markdown("### Generated Questions:")
        questions_list = questions.split("\n\n")
        for q in questions_list:
            if question_type == "Multiple Choice":
                display_multiple_choice(q, show_answers)
            else:
                st.markdown(q)

# Function to display multiple-choice questions with conditional answer display
def display_multiple_choice(question_text, show_answers):
    parts = question_text.split("\n")
    if len(parts) >= 5:
        st.markdown(f"**{parts[0]}**")
        for part in parts[1:5]:
            st.markdown(f"- {part}")
        # Conditionally display the answer
        if show_answers:
            st.markdown(f"**Answer:** {parts[5]}")
    else:
        st.markdown(question_text)

# Function to generate DOCX file
def generate_docx(questions, question_type, show_answers):
    doc = Document()
    doc.add_heading("Generated Questions", 0)

    questions_list = questions.split("\n\n")
    for q in questions_list:
        if question_type == "Multiple Choice":
            parts = q.split("\n")
            if len(parts) >= 5:
                doc.add_paragraph(parts[0], style='List Bullet')
                for part in parts[1:5]:
                    doc.add_paragraph(part, style='List Bullet 2')
                if show_answers:
                    doc.add_paragraph(f"Answer: {parts[5]}")
            else:
                doc.add_paragraph(q)
        else:
            doc.add_paragraph(q)

    return doc

# Function to generate PDF file
def generate_pdf(questions, question_type, show_answers):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    y = 750
    for q in questions.split("\n\n"):
        if question_type == "Multiple Choice":
            parts = q.split("\n")
            if len(parts) >= 5:
                can.drawString(50, y, parts[0])
                y -= 20
                for part in parts[1:5]:
                    can.drawString(70, y, part)
                    y -= 20
                if show_answers:
                    can.drawString(50, y, f"Answer: {parts[5]}")
                    y -= 20
            else:
                can.drawString(50, y, q)
                y -= 20
        else:
            can.drawString(50, y, q)
            y -= 30
        y -= 20  # Add extra space between questions
        if y < 50:  # Start a new page if needed
            can.showPage()
            y = 750
    can.save()

    packet.seek(0)
    new_pdf = PdfReader(packet)
    output = PdfWriter()
    output.add_page(new_pdf.pages[0])
    return output

# Initialize session state for storing recent generated questions
if "recent_generated" not in st.session_state:
    st.session_state.recent_generated = []

# Initialize session state for API key and login status
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

# Streamlit UI setup
st.set_page_config(page_title="AI Question Generator?", page_logo="logo1.png", layout="wide") # Updated title and icon

# Custom CSS for enhanced UI (you can add more styles here)
st.markdown("""
<style>
body {
    color: #4f4f4f;
    background-color: #f8f9fa;
}
.stButton>button {
    color: #ffffff;
    background-color: #007bff;
    border: none;
    border-radius: 5px;
    padding: 10px 20px;
}
.stTextInput>div>div>input {
    border: 2px solid #ced4da;
    border-radius: 5px;
}
.stSlider>div>div>div>div {
    background: #007bff;
}
</style>
""", unsafe_allow_html=True)

st.title("❓Question") # Updated title

# --- Login ---
if not st.session_state.logged_in:
    api_key_input = st.text_input("Enter your Gemini API key:", type="password")
    if st.button("Login"):
        if is_valid_api_key(api_key_input):
            st.session_state.api_key = api_key_input
            st.session_state.logged_in = True
            st.success("Login successful!")
            st.rerun()  # Rerun the app to show the main UI
        else:
            st.error("Invalid API key. Please try again.")
else:
    # --- Main UI ---
    model = initialize_model(st.session_state.api_key)

    # Configuration sidebar
    with st.sidebar:
        st.header("Question Settings")
        question_type = st.selectbox("Select Question Type", ["Multiple Choice", "Fill in the Blank", "True or False", "Open-ended"])
        num_questions = st.slider("Number of Questions", 1, 20, 5)

        show_answers = st.checkbox("Show Answers", value=False)

    # Recent Generated sidebar with toggle
    with st.sidebar:
        if st.checkbox("Show Recent Generated"):
            st.header("Recent Generated")
            for question in st.session_state.recent_generated:
                st.markdown(question)

    # Main area for topic input and results
    topic = st.text_input("Enter the topic for question generation", placeholder="e.g., Photosynthesis, World War II")

    if st.button("Generate Questions"):
        if topic:
            with st.spinner('Generating questions...'):
                generated_questions = generate_questions(model, topic, num_questions, question_type, show_answers)
                display_questions(generated_questions, question_type, show_answers)

                # Store recent generated questions in session state
                if generated_questions:
                    questions_list = generated_questions.split("\n\n")
                    # Add new questions at the beginning of the list (no limit on the number of questions)
                    st.session_state.recent_generated = questions_list + st.session_state.recent_generated

                # Download buttons
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    if generated_questions:
                        docx_doc = generate_docx(generated_questions, question_type, show_answers)
                        docx_buffer = BytesIO()
                        docx_doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        st.download_button(
                            label="Download as DOCX",
                            data=docx_buffer,
                            file_name="generated_questions.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                with col2:
                    if generated_questions:
                        pdf_output = generate_pdf(generated_questions, question_type, show_answers)
                        pdf_buffer = BytesIO()
                        pdf_output.write(pdf_buffer)
                        pdf_buffer.seek(0)
                        st.download_button(
                            label="Download as PDF",
                            data=pdf_buffer,
                            file_name="generated_questions.pdf",
                            mime="application/pdf"
                        )
        else:
            st.error("Please enter a topic.")