import streamlit as st
import google.generativeai as genai  # Ensure this is installed and up-to-date
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# ---------------------------------------------------------
# Configure the Google Generative AI API
# ---------------------------------------------------------
genai.configure(api_key="AIzaSyC8zKSK8_xurGABNkBGyn-bbVj4mho-5B8")  # Replace with your actual API key

generation_config = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
}

# Initialize the generative model (adjust model_name if needed)
model = genai.GenerativeModel(model_name="gemini-2.0-flash", generation_config=generation_config)

# ---------------------------------------------------------
# API-based Generation Functions
# ---------------------------------------------------------
def generate_questions(topic, num_questions, qtype):
    if qtype == "Multiple Choice":
        prompt = (f"Generate {num_questions} multiple-choice questions on the topic '{topic}'. "
                  f"Each question should include 4 options (A, B, C, D) and indicate the correct answer.")
    elif qtype == "True or False":
        prompt = (f"Generate {num_questions} true or false questions on the topic '{topic}'. "
                  f"Include the correct answer for each question.")
    elif qtype == "Fill in the Blanks":
        prompt = (f"Generate {num_questions} fill in the blank questions on the topic '{topic}'. "
                  f"Provide the answer for each blank.")
    else:
        prompt = f"Generate {num_questions} questions on the topic '{topic}'."
    try:
        with st.spinner("Generating questions..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating questions: {e}")
        return f"Error generating questions: {e}"

def generate_titles(topic, num_titles):
    prompt = f"Generate {num_titles} research-worthy thesis or capstone project titles for the topic '{topic}'."
    try:
        with st.spinner("Generating titles..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating titles: {e}")
        return f"Error generating titles: {e}"

def generate_reviewer(content, title):
    prompt = f"Generate study materials and key points for the title '{title}' based on the following content: {content}"
    try:
        with st.spinner("Generating reviewer content..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating reviewer content: {e}")
        return f"Error generating reviewer content: {e}"

def generate_essay_summary(text):
    prompt = f"Provide a clear and concise summary of the following text:\n\n{text}"
    try:
        with st.spinner("Generating essay summary..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating essay summary: {e}")
        return f"Error generating essay summary: {e}"

def generate_citations(style, source):
    prompt = f"Generate a citation for the source '{source}' in {style} format."
    try:
        with st.spinner("Generating citation..."):
            response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating citation: {e}")
        return f"Error generating citation: {e}"

# ---------------------------------------------------------
# Export Functions
# ---------------------------------------------------------
def export_docx(content):
    doc = Document()
    doc.add_heading("Generated Content", level=0)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_pdf(content):
    buffer = BytesIO()
    can = canvas.Canvas(buffer, pagesize=letter)
    text_object = can.beginText(40, 750)
    for line in content.split('\n'):
        text_object.textLine(line)
    can.drawText(text_object)
    can.showPage()
    can.save()
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# Initialize Session State
# ---------------------------------------------------------
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "generated_content" not in st.session_state:
    st.session_state.generated_content = ""
if "dashboard" not in st.session_state:
    st.session_state.dashboard = []  # List of tuples: (content type, content)

# ---------------------------------------------------------
# App Title & Tagline
# ---------------------------------------------------------
st.title("EduGen: Smart Academic Generator")
st.markdown("*Making Learning & Teaching Easier* :mortar_board:")
st.markdown("---")

# ---------------------------------------------------------
# User Login Section
# ---------------------------------------------------------
if not st.session_state.logged_in:
    with st.container():
        st.header("User Login")
        col1, col2 = st.columns([1,3]) # Adjust column ratio as needed
        with col1:
            username = st.text_input("Username", key="username_login")
        with col2:
            password = st.text_input("Password", type="password", key="password_login")
        if st.button("Login", key="login_button"):
            # Dummy authentication; replace with real authentication as needed
            if username == "admin" and password == "password":
                st.session_state.logged_in = True
                st.success("Logged in successfully!", icon="✅")
                st.rerun() # Rerun to update UI after login
            else:
                st.error("Invalid credentials. Please try again.", icon="🚨")
    st.stop()  # Stop further execution until logged in

# ---------------------------------------------------------
# Main App Content after Login
# ---------------------------------------------------------
st.sidebar.title("Dashboard")
generator_option = st.sidebar.radio("Select Generator", (
    "Question Generator",
    "Title Generator",
    "Reviewer Creator",
    "Essay & Summary Generator",
    "Citation & Bibliography Tool"
))

# Initialize content variable outside the conditional blocks
content = ""

if generator_option == "Question Generator":
    st.header("Question Generator")
    topic = st.text_input("Enter the topic for questions", placeholder="e.g., Photosynthesis")
    num_questions = st.number_input("Number of Questions", min_value=1, max_value=20, value=5)
    qtype = st.selectbox("Select Question Type", ("Multiple Choice", "True or False", "Fill in the Blanks"))
    if st.button("Generate Questions"):
        content = generate_questions(topic, num_questions, qtype)
        st.session_state.generated_content = content
        st.session_state.dashboard.append(("Questions", content))

elif generator_option == "Title Generator":
    st.header("Thesis & Capstone Title Generator")
    topic = st.text_input("Enter the topic for titles", placeholder="e.g., Impact of Social Media on Teenagers")
    num_titles = st.number_input("Number of Titles", min_value=1, max_value=10, value=3)
    if st.button("Generate Titles"):
        content = generate_titles(topic, num_titles)
        st.session_state.generated_content = content
        st.session_state.dashboard.append(("Titles", content))

elif generator_option == "Reviewer Creator":
    st.header("Reviewer Creator")
    title_input = st.text_input("Enter the title for the study material", placeholder="e.g., Chapter 3 - Cell Biology")
    content_input = st.text_area("Enter content or topics", height=150, placeholder="Key concepts, formulas, etc.")
    if st.button("Generate Reviewer"):
        content = generate_reviewer(content_input, title_input)
        st.session_state.generated_content = content
        st.session_state.dashboard.append(("Reviewer", content))

elif generator_option == "Essay & Summary Generator":
    st.header("Essay & Summary Generator")
    essay_input = st.text_area("Enter your essay prompt or content", height=150, placeholder="Paste essay text or type a prompt...")
    if st.button("Generate Summary"):
        content = generate_essay_summary(essay_input)
        st.session_state.generated_content = content
        st.session_state.dashboard.append(("Essay Summary", content))

elif generator_option == "Citation & Bibliography Tool":
    st.header("Citation & Bibliography Tool")
    source = st.text_input("Enter the source title or URL", placeholder="e.g., www.example.com or 'The Great Gatsby'")
    style = st.selectbox("Select Citation Style", ("APA", "MLA", "Chicago"))
    if st.button("Generate Citation"):
        content = generate_citations(style, source)
        st.session_state.generated_content = content
        st.session_state.dashboard.append(("Citation", content))

# ---------------------------------------------------------
# Display Generated Content and Edit Area
# ---------------------------------------------------------
if st.session_state.generated_content:
    st.header("Generated Content")
    content = st.text_area("Generated Output", value=st.session_state.generated_content, height=300)
    st.session_state.generated_content = content # Update session state with potentially edited content

# ---------------------------------------------------------
# Export Options: Download as DOCX or PDF
# ---------------------------------------------------------
if st.session_state.generated_content:
    st.header("Export Options")
    col1, col2 = st.columns(2)
    with col1:
        docx_buffer = export_docx(st.session_state.generated_content)
        st.download_button(
            label="Download as DOCX",
            data=docx_buffer,
            file_name="generated_content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col2:
        pdf_buffer = export_pdf(st.session_state.generated_content)
        st.download_button(
            label="Download as PDF",
            data=pdf_buffer,
            file_name="generated_content.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    st.markdown("---") # Separator after export options

# ---------------------------------------------------------
# Sidebar: Recently Generated Content Overview - Enhanced
# ---------------------------------------------------------
st.sidebar.subheader("Recent Content")
if st.session_state.dashboard:
    for gen_type, gen_content in st.session_state.dashboard:
        with st.sidebar.expander(f"**{gen_type}:** {gen_content[:50]}...", expanded=False): # Expandable recent content
            st.write(gen_content)
else:
    st.sidebar.write("No content generated yet.")